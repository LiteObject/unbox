"""Word (.docx) text extractor using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from unbox.base import BaseExtractor


class DocxExtractor(BaseExtractor):
    """Extract plain text from Word documents."""

    supported_extensions = [".docx"]

    def extract(self, file_path: Path) -> str:
        """Extract text from paragraphs and tables of a Word document.

        Parameters
        ----------
        file_path:
            Path to the ``.docx`` file.

        Returns
        -------
        str
            Paragraphs joined by newlines, followed by table content.
        """
        doc = Document(str(file_path))
        parts: list[str] = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows_text: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append("\n".join(rows_text))

        return "\n\n".join(parts)
